"""
使用 python-docx 将 doc\\minfadian.docx 转换为纯文本
"""

from pathlib import Path
from docx import Document


def extract_text_from_docx(docx_path: str, output_path: str | None = None) -> str:
    """
    从 docx 文件中提取纯文本
    
    Args:
        docx_path: docx 文件路径
        output_path: 可选的输出文件路径
    
    Returns:
        提取的纯文本内容
    """
    doc = Document(docx_path)
    
    # 提取所有段落的文本
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    
    # 提取表格中的文本（如果有）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)
    
    full_text = "\n".join(paragraphs)
    
    # 如果指定了输出路径，写入文件
    if output_path:
        Path(output_path).write_text(full_text, encoding="utf-8")
        print(f"文本已保存到: {output_path}")
    
    return full_text


if __name__ == "__main__":
    docx_file = r"c:\Users\echuzhi\repo\legal-agent\doc\xianfa.docx"
    output_file = r"c:\Users\echuzhi\repo\legal-agent\doc\xianfa.txt"
    
    text = extract_text_from_docx(docx_file, output_file)
    print(f"共提取 {len(text)} 字符的文本")
    print("\n前 500 字符预览:")
    print(text[:500])