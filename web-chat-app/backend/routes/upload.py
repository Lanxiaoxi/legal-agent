"""文件上传路由模块"""
import asyncio
import json
import logging
import shutil
import uuid
from datetime import datetime, timezone
from pathlib import Path

import docx
import fitz  # pymupdf
from fastapi import APIRouter, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse

from config import config
from services.pdf_routing import extract_pdf_text, try_ocr_page

logger = logging.getLogger(__name__)

router = APIRouter()

# 支持的文件类型
ALLOWED_EXTENSIONS = {".txt", ".pdf", ".docx", ".jpg", ".jpeg", ".png", ".bmp"}
# 单次会话最大总文件大小
MAX_SESSION_SIZE_MB = 50

# ---------- 文本提取 ----------


def _extract_image_text(file_path: Path) -> str:
    """从图片中提取文字（通过 OCR）"""
    # 图片没有文字层，直接 OCR。
    # 注意：PyMuPDF 的 get_textpage_ocr 只支持 PDF 文档，图片需先转为单页 PDF 再识别，
    # 否则会报 "source or target not a PDF"。
    img_doc = fitz.open(str(file_path))  # 图片按单页文档打开
    try:
        pix = img_doc[0].get_pixmap()
    finally:
        img_doc.close()

    pdf_doc = fitz.open()  # 空的内存 PDF
    try:
        page = pdf_doc.new_page(width=pix.width, height=pix.height)
        page.insert_image(fitz.Rect(0, 0, pix.width, pix.height), pixmap=pix)
        text = try_ocr_page(page, 1)
    finally:
        pdf_doc.close()

    if not text:
        raise RuntimeError(
            f"无法从图片中识别出文字。"
            "请确认图片包含清晰的文字内容，或已安装 Tesseract OCR。"
        )
    return text


def _extract_text(file_path: Path, ext: str) -> tuple[str, dict]:
    """从文件中提取纯文本，返回 (text, extra)

    extra: PDF 时包含 route_summary / routed 信息，其余类型为空 dict
    """
    if ext == ".txt":
        text = file_path.read_text(encoding="utf-8", errors="replace")
        if not text.strip():
            # 尝试以其他编码读取
            for enc in ["gbk", "gb2312", "latin-1"]:
                try:
                    text = file_path.read_text(encoding=enc)
                    if text.strip():
                        logger.info(f"File {file_path.name} decoded with {enc}")
                        break
                except (UnicodeDecodeError, LookupError):
                    continue
        return text, {}

    if ext == ".docx":
        doc = docx.Document(str(file_path))
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        # 也提取表格内容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        paragraphs.append(text)
        return "\n".join(paragraphs), {}

    if ext in {".jpg", ".jpeg", ".png", ".bmp"}:
        return _extract_image_text(file_path), {}

    if ext == ".pdf":
        # 四路页面级分流（pdf-inspector），失败/禁用时自动回退 PyMuPDF + OCR
        result = extract_pdf_text(file_path, config.use_pdf_routing)
        return result.text, {"route_summary": result.route_summary, "routed": result.routed}

    return "", {}


# ---------- 分块 ----------


def _chunk_text(text: str, chunk_size: int = 2000, overlap: int = 100) -> list[str]:
    """将文本按自然段落分块，每块不超过 chunk_size 字"""
    paragraphs = text.split("\n")
    chunks = []
    current = ""
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        if len(current) + len(para) + 1 <= chunk_size:
            current = (current + "\n" + para).strip() if current else para
        else:
            if current:
                chunks.append(current)
            # 如果单个段落超过 chunk_size，硬切
            if len(para) > chunk_size:
                for i in range(0, len(para), chunk_size - overlap):
                    chunks.append(para[i:i + chunk_size])
            else:
                current = para
    if current:
        chunks.append(current)
    return chunks


# ---------- 存储管理 ----------


def _get_session_dir(session_id: str) -> Path:
    return Path(config.upload_dir) / session_id


def _get_metadata_path(session_id: str) -> Path:
    return _get_session_dir(session_id) / "metadata.json"


def _get_chunks_dir(session_id: str) -> Path:
    return _get_session_dir(session_id) / "chunks"


def _load_metadata(session_id: str) -> dict:
    path = _get_metadata_path(session_id)
    if path.exists():
        return json.loads(path.read_text(encoding="utf-8"))
    return {"files": [], "last_accessed_at": datetime.now(timezone.utc).isoformat()}


def _save_metadata(session_id: str, metadata: dict):
    path = _get_metadata_path(session_id)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(metadata, ensure_ascii=False, indent=2), encoding="utf-8")


def _touch_access(session_id: str):
    """更新 last_accessed_at"""
    metadata = _load_metadata(session_id)
    metadata["last_accessed_at"] = datetime.now(timezone.utc).isoformat()
    _save_metadata(session_id, metadata)


def _get_session_total_size(session_id: str) -> int:
    """获取 session 总存储大小（字节）"""
    session_dir = _get_session_dir(session_id)
    if not session_dir.exists():
        return 0
    total = 0
    for f in session_dir.rglob("*"):
        if f.is_file():
            total += f.stat().st_size
    return total


# ---------- API 端点 ----------


@router.post("/api/upload")
async def upload_file(
    file: UploadFile = File(...),
    session_id: str = Form(default=""),
    user_id: str = Form(default=""),
):
    """上传文件并登记，文本提取/分块由后台任务异步完成

    上传接口只负责：校验 → 保存原文件 → 登记 metadata（status=processing）→ 立即返回。
    文本提取与分块放入后台线程池任务（asyncio.to_thread），不阻塞事件循环；
    处理完成后更新状态为 done，失败为 failed（保留原文件供重试）。

    Args:
        file: 上传的文件
        session_id: 会话 ID
        user_id: 匿名用户 ID

    Returns:
        {files: [{id, name, type, size, status, ...}]}
    """
    if not session_id:
        raise HTTPException(status_code=400, detail="session_id is required")

    # 验证文件扩展名
    filename = file.filename or "unknown"
    ext = Path(filename).suffix.lower()

    # .doc 旧版格式不再支持，引导用户转换
    if ext == ".doc":
        raise HTTPException(
            status_code=400,
            detail="暂不支持 .doc 格式，请先用 Word / WPS 将文档另存为 .docx 后再上传。"
        )

    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件类型 '{ext}'，仅支持: {', '.join(ALLOWED_EXTENSIONS)}"
        )

    # 验证大小
    file.file.seek(0, 2)
    file_size = file.file.tell()
    file.file.seek(0)
    max_bytes = config.upload_max_size_mb * 1024 * 1024
    if file_size > max_bytes:
        raise HTTPException(
            status_code=413,
            detail=f"文件大小 {file_size / 1024 / 1024:.1f}MB 超过限制 {config.upload_max_size_mb}MB"
        )

    # 检查 session 总文件大小
    current_total = _get_session_total_size(session_id)
    if current_total + file_size > MAX_SESSION_SIZE_MB * 1024 * 1024:
        raise HTTPException(
            status_code=413,
            detail=f"会话文件总大小超过 {MAX_SESSION_SIZE_MB}MB 限制"
        )

    # 生成 file_id，保存原始文件（后台处理成功后删除；失败保留供重试）
    file_id = str(uuid.uuid4())[:8]
    session_dir = _get_session_dir(session_id)
    chunks_dir = _get_chunks_dir(session_id)
    session_dir.mkdir(parents=True, exist_ok=True)
    chunks_dir.mkdir(parents=True, exist_ok=True)

    saved_path = session_dir / f"{file_id}{ext}"
    content = await file.read()
    saved_path.write_bytes(content)

    # 登记 metadata，状态 processing
    metadata = _load_metadata(session_id)
    file_info = {
        "id": file_id,
        "name": filename,
        "type": ext.lstrip("."),
        "size": file_size,
        "char_count": 0,
        "chunk_count": 0,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "status": "processing",
        "error": None,
    }
    metadata["files"].append(file_info)
    metadata["user_id"] = user_id or None
    metadata["last_accessed_at"] = datetime.now(timezone.utc).isoformat()
    _save_metadata(session_id, metadata)

    # 后台异步处理（提取 + 分块），不阻塞上传请求
    _schedule_process(session_id, file_id, saved_path, ext)

    logger.info(
        f"[UPLOAD] session={session_id} user={user_id} file={filename} "
        f"size={file_size}B file_id={file_id} status=processing (async)"
    )

    return {"files": metadata["files"]}


@router.get("/api/files/{session_id}")
async def list_files(session_id: str):
    """列出某 session 的上传文件

    Args:
        session_id: 会话 ID

    Returns:
        {files: [...]}
    """
    _touch_access(session_id)
    metadata = _load_metadata(session_id)
    if not metadata.get("files"):
        return {"files": []}
    return {"files": metadata["files"]}


@router.get("/api/files/{session_id}/{file_id}")
async def download_file(session_id: str, file_id: str):
    """下载某 session 的指定文件（目前仅支持 AI 生成的 .docx 文档）

    Args:
        session_id: 会话 ID
        file_id: 文件 ID

    Returns:
        文件下载响应
    """
    # 查找 AI 生成的文档
    files_dir = _get_session_dir(session_id) / "files"
    docx_path = files_dir / f"{file_id}.docx"
    if not docx_path.exists():
        raise HTTPException(status_code=404, detail="文件不存在或已过期")

    metadata = _load_metadata(session_id)
    file_info = next((f for f in metadata.get("files", []) if f.get("id") == file_id), None)
    download_name = file_info.get("name", f"{file_id}.docx") if file_info else f"{file_id}.docx"

    return FileResponse(
        path=str(docx_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=download_name,
    )


@router.delete("/api/files/{session_id}")
async def delete_files(session_id: str):
    """清理某 session 的所有上传文件

    Args:
        session_id: 会话 ID
    """
    session_dir = _get_session_dir(session_id)
    if session_dir.exists():
        shutil.rmtree(session_dir)
        logger.info(f"[CLEANUP] Deleted files for session={session_id}")
    return {"status": "ok", "message": "Files cleaned up"}


@router.delete("/api/files/{session_id}/{file_id}")
async def delete_file(session_id: str, file_id: str):
    """删除某 session 的指定文件（上传的文本分块或 AI 生成的文档）

    Args:
        session_id: 会话 ID
        file_id: 文件 ID

    Returns:
        {files: 删除后剩余的该会话文件列表}
    """
    metadata = _load_metadata(session_id)
    files = metadata.get("files", [])
    file_info = next((f for f in files if f.get("id") == file_id), None)
    if not file_info:
        raise HTTPException(status_code=404, detail="文件不存在或已过期")

    session_dir = _get_session_dir(session_id)
    if file_info.get("generated"):
        # AI 生成的 Word 文档
        docx_path = session_dir / "files" / f"{file_id}.docx"
        if docx_path.exists():
            docx_path.unlink(missing_ok=True)
    else:
        # 上传文件的分块
        chunks_dir = _get_chunks_dir(session_id)
        if chunks_dir.exists():
            for chunk_path in chunks_dir.glob(f"{file_id}_*.txt"):
                chunk_path.unlink(missing_ok=True)
        # 原始文件（processing/failed 期间保留，用于后台处理/重试）
        orig_path = session_dir / f"{file_id}.{file_info.get('type', '')}"
        orig_path.unlink(missing_ok=True)

    # 从元数据移除
    metadata["files"] = [f for f in files if f.get("id") != file_id]
    metadata["last_accessed_at"] = datetime.now(timezone.utc).isoformat()
    _save_metadata(session_id, metadata)

    logger.info(f"[DELETE] session={session_id} file={file_id} name={file_info.get('name')}")
    return {"files": metadata["files"]}


# ---------- 异步后台处理（上传与处理分离） ----------

# 持有后台任务引用，防止被 GC；任务完成自动移除
_background_tasks: set = set()


def _file_record_exists(session_id: str, file_id: str) -> bool:
    """检查文件记录是否仍存在（可能已被用户删除）"""
    metadata = _load_metadata(session_id)
    return any(f.get("id") == file_id for f in metadata.get("files", []))


def _update_file_status(session_id: str, file_id: str, status: str, **fields) -> bool:
    """更新某文件的状态字段，返回是否找到记录"""
    metadata = _load_metadata(session_id)
    for f in metadata.get("files", []):
        if f.get("id") == file_id:
            f["status"] = status
            f.update(fields)
            break
    else:
        return False  # 记录不存在（文件已被删除）
    metadata["last_accessed_at"] = datetime.now(timezone.utc).isoformat()
    _save_metadata(session_id, metadata)
    return True


def _process_file_sync(session_id: str, file_id: str, file_path: Path, ext: str) -> None:
    """同步执行 提取文本 + 分块落盘 + 更新状态（在后台线程池中运行）

    成功 → status=done，删除原文件只留分块
    失败 → status=failed，保留原文件供前端重试
    """
    # 处理期间文件可能已被用户删除，跳过避免写回无效记录
    if not _file_record_exists(session_id, file_id):
        logger.info(f"[PROCESS] file {file_id} removed during processing, skip")
        return

    try:
        text, extract_extra = _extract_text(file_path, ext)
        if not text.strip():
            detail = (
                f"无法从文件中提取文本内容（{file_path.name}，{file_path.stat().st_size} 字节）。"
                f"{'TXT 文件可能是空文件或编码不受支持。' if ext == '.txt' else ''}"
                f"{'DOCX 文件可能只包含图片或格式不受支持。' if ext == '.docx' else ''}"
            )
            _update_file_status(session_id, file_id, "failed", error=detail)
            return

        chunks = _chunk_text(text)
        chunks_dir = _get_chunks_dir(session_id)
        chunks_dir.mkdir(parents=True, exist_ok=True)
        for i, chunk in enumerate(chunks, 1):
            (chunks_dir / f"{file_id}_{i:03d}.txt").write_text(chunk, encoding="utf-8")

        _update_file_status(
            session_id, file_id, "done",
            char_count=len(text),
            chunk_count=len(chunks),
            extract_mode="pdf_routed" if extract_extra.get("routed") else "legacy",
            route_summary=extract_extra.get("route_summary"),
            error=None,
        )
        # 处理成功，删除原文件只留分块
        file_path.unlink(missing_ok=True)
        logger.info(
            f"[PROCESS] session={session_id} file={file_id} done, "
            f"chunks={len(chunks)}, removed original"
        )
    except Exception as e:
        logger.error(f"[PROCESS] session={session_id} file={file_id} failed: {e}")
        # 失败保留原文件供重试
        _update_file_status(session_id, file_id, "failed", error=str(e))


def _schedule_process(session_id: str, file_id: str, file_path: Path, ext: str) -> None:
    """将文件处理任务放入后台（线程池执行，不阻塞事件循环）"""

    async def _run():
        try:
            await asyncio.to_thread(_process_file_sync, session_id, file_id, file_path, ext)
        except Exception as e:
            logger.error(f"[PROCESS] unexpected error session={session_id} file={file_id}: {e}")
            _update_file_status(session_id, file_id, "failed", error=f"处理失败: {e}")

    task = asyncio.create_task(_run())
    _background_tasks.add(task)
    task.add_done_callback(_background_tasks.discard)


@router.post("/api/files/{session_id}/{file_id}/retry")
async def retry_file(session_id: str, file_id: str):
    """重试失败的文件处理（不重新上传，直接对保留的原文件重新提取）

    Args:
        session_id: 会话 ID
        file_id: 文件 ID

    Returns:
        {files: 该会话文件列表}
    """
    metadata = _load_metadata(session_id)
    file_info = next((f for f in metadata.get("files", []) if f.get("id") == file_id), None)
    if not file_info:
        raise HTTPException(status_code=404, detail="文件不存在或已过期")

    status = file_info.get("status")
    if status == "processing":
        return {"files": metadata["files"], "message": "文件正在处理中，请稍候"}
    if status == "done":
        return {"files": metadata["files"], "message": "文件已处理完成"}

    # failed → 重新处理
    file_type = file_info.get("type", "")
    orig_path = _get_session_dir(session_id) / f"{file_id}.{file_type}"
    if not orig_path.exists():
        raise HTTPException(
            status_code=400,
            detail="原始文件已丢失，无法重试，请重新上传"
        )

    _update_file_status(session_id, file_id, "processing", error=None)
    _schedule_process(session_id, file_id, orig_path, f".{file_type}")

    logger.info(f"[RETRY] session={session_id} file={file_id} rescheduled")
    return {"files": _load_metadata(session_id)["files"], "message": "已开始重新处理"}


# ---------- TTL 清理 ----------


async def cleanup_expired_files():
    """清理超过 TTL 的 session 目录（后台任务，由 main.py 启动）"""
    import asyncio

    while True:
        await asyncio.sleep(3600)  # 每小时检查一次
        upload_dir = Path(config.upload_dir)
        if not upload_dir.exists():
            continue

        ttl_seconds = config.upload_ttl_days * 24 * 3600
        now = datetime.now(timezone.utc)

        for session_dir in upload_dir.iterdir():
            if not session_dir.is_dir():
                continue
            metadata_path = session_dir / "metadata.json"
            if not metadata_path.exists():
                continue
            try:
                metadata = json.loads(metadata_path.read_text(encoding="utf-8"))
                last_accessed = datetime.fromisoformat(
                    metadata.get("last_accessed_at", "2000-01-01T00:00:00+00:00")
                )
                if (now - last_accessed).total_seconds() > ttl_seconds:
                    shutil.rmtree(session_dir)
                    logger.info(
                        f"[TTL] Cleaned up expired session={session_dir.name}, "
                        f"last_accessed={last_accessed.isoformat()}"
                    )
            except Exception as e:
                logger.warning(f"[TTL] Error processing {session_dir}: {e}")
