"""
使用 python-docx 将 doc\\minfadian.docx 转换为纯文本
"""

from pathlib import Path
from docx import Document


def extract_text_from_docx(docx_path: str, output_path: str | None = None) -> str:
    """
    从 docx 文件中提取纯文本
    
    Args:
        docx_path: docx 文件路径
        output_path: 可选的输出文件路径
    
    Returns:
        提取的纯文本内容
    """
    doc = Document(docx_path)
    
    # 提取所有段落的文本
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    
    # 提取表格中的文本（如果有）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)
    
    full_text = "\n".join(paragraphs)
    
    # 如果指定了输出路径，写入文件
    if output_path:
        Path(output_path).write_text(full_text, encoding="utf-8")
        print(f"文本已保存到: {output_path}")
    
    return full_text


if __name__ == "__main__":
    import glob
    
    doc_dir = r"c:\Users\echuzhi\repo\legal-agent\doc"
    docx_files = glob.glob(f"{doc_dir}/*.docx")
    
    print(f"找到 {len(docx_files)} 个 docx 文件\n")
    
    for docx_file in sorted(docx_files):
        docx_path = Path(docx_file)
        output_file = str(docx_path.with_suffix('.txt'))
        
        print(f"处理: {docx_path.name}")
        text = extract_text_from_docx(docx_file, output_file)
        print(f"  -> 已保存到: {docx_path.stem}.txt ({len(text)} 字符)\n")